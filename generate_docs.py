import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_documentation():
    doc = Document()

    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    # Primary: #0F172A (Deep Slate), Accent: #059669 (Emerald), Secondary: #0284C7 (Blue)
    COLOR_PRIMARY = RGBColor(15, 23, 42)
    COLOR_EMERALD = RGBColor(5, 150, 105)
    COLOR_MUTED = RGBColor(100, 116, 139)

    # Document Header / Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("NexFin — Platform Technical & Feature Documentation")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Next-Generation AI Financial Digital Twin, Liquidity Telemetry & MSME Solvency Platform")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = COLOR_EMERALD
    run_sub.font.bold = True

    # Metadata box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Platform Name:", "NexFin (formerly FinTwin)"), ("Release Version:", "v2.5 Enterprise Production")],
        [("Target Audience:", "MSMEs, CFOs, Founders, Accountants"), ("Documentation Date:", "September 2026")]
    ]
    for r_idx, row in enumerate(meta_table.rows):
        for c_idx, cell in enumerate(row.cells):
            label, val = meta_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(label + " ")
            r1.font.bold = True
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = COLOR_PRIMARY
            r2 = p.add_run(val)
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = COLOR_MUTED
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, 80, 80, 120, 120)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for Section Headings
    def add_section_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = COLOR_PRIMARY
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_EMERALD
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_bullet(p_or_text, bold_prefix="", text=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r1 = p.add_run(bold_prefix + ": ")
            r1.font.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = COLOR_PRIMARY
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(51, 65, 85)

    # -------------------------------------------------------------
    # 1. EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    add_section_heading("1. Executive Summary & Value Proposition")
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(8)
    p1.add_run(
        "NexFin is an autonomous AI-powered Financial Digital Twin built specifically for Indian Micro, Small, and Medium Enterprises (MSMEs) and fast-growing enterprises. "
        "It acts as a 24/7 predictive co-pilot that ingests multi-source financial telemetry (consented GST e-invoices, bank feeds, expense ledgers, and payroll streams) "
        "to simulate cash runway, forecast debtor payment delays with trained machine learning models, enforce MSMED Act statutory 45-day recovery timelines, and optimize tax liabilities."
    )

    # -------------------------------------------------------------
    # 2. COMPLETE TECHNOLOGY STACK
    # -------------------------------------------------------------
    add_section_heading("2. Full Technology Stack Architecture")

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = "Tier / Component"
    hdr_cells[1].text = "Technology / Libraries"
    hdr_cells[2].text = "Architectural Purpose"
    for cell in hdr_cells:
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 100, 100, 120, 120)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)

    tech_rows = [
        ("Frontend Framework", "React 18 / 19, Vite 8.2", "Ultra-fast Single Page Application with Hot Module Replacement"),
        ("Styling & Design System", "Vanilla CSS3, CSS Custom Properties", "Tailored executive fintech theme, glassmorphism, responsive grid"),
        ("Data Visualization", "Recharts (v2.x)", "Interactive AreaCharts, BarCharts, Donut gauges & predictive bounds"),
        ("Icons & UI Graphics", "Lucide React, Custom SVG Vector Suite", "Unified modern iconography and dynamic responsive logo mark"),
        ("Routing & Navigation", "React Router DOM v6", "Client-side declarative routing, deep-linking & role-based redirects"),
        ("Multilingual Localization", "React Context API (18 Indian Languages)", "Dynamic on-the-fly translation (Hindi, Marathi, Gujarati, Tamil, etc.)"),
        ("File & Data Ingestion", "PapaParse (CSV), XLSX / SheetJS", "Client-side multi-format parsing for CSV, Excel, PDF scans, and JSON"),
        ("Backend Framework", "Python 3.10+, FastAPI, Uvicorn", "High-throughput asynchronous REST API & ML microservices"),
        ("Data Modeling & Schemas", "Pydantic v2", "Strict request/response validation and data serialization"),
        ("Machine Learning Engine", "Scikit-Learn, LightGBM", "Dual-engine trained on 88,305 B2B invoices for delay & risk prediction"),
        ("Deployment & Hosting", "Vercel (Frontend), Render / AWS (Backend)", "Global Edge CDN deployment with automated GitHub CI/CD webhooks")
    ]

    for tier, tech, purpose in tech_rows:
        row = tech_table.add_row()
        for idx, text in enumerate([tier, tech, purpose]):
            cell = row.cells[idx]
            cell.text = text
            set_cell_margins(cell, 70, 70, 100, 100)
            set_cell_background(cell, "F8FAFC" if len(tech_table.rows) % 2 == 0 else "FFFFFF")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # 3. CORE MODULE FEATURES
    # -------------------------------------------------------------
    add_section_heading("3. Comprehensive Module Feature Breakdown")

    modules = [
        ("3.1 Executive Dashboard", [
            ("Real-Time Solvency Telemetry", "Instant visibility into Opening Liquid Cash, Total Receivables, Monthly Burn Rate, and Net Runway Buffer."),
            ("GST & Financial Health Monitor", "Real-time computed composite health index (0–100 score) evaluating Runway Solvency, Receivables Aging, Concentration Risk, and GST Reserves with animated SVG radial gauge."),
            ("Autonomous Runway Gauge", "Color-coded dynamic progress tracker calculating operational runway in days (Healthy/Moderate/Critical Burn)."),
            ("AI Cash Velocity & Confidence Radar", "Forecast projections across 30, 60, and 90-day time horizons with historical reference lines."),
            ("Receivables Aging Matrix", "Categorized breakdown into 0–30 Days, 31–60 Days, 61–90 Days, and 90+ Days buckets with risk indicators.")
        ]),
        ("3.2 Cash Flow Digital Twin", [
            ("90-Day Predictive Timeline", "Granular day-by-day cash balance projection integrating scheduled receivables and recurring liabilities."),
            ("Monte Carlo Scenario Bands", "Visual confidence bounds comparing Conservative (stress), Base Case, and Optimistic liquidity trajectories."),
            ("Daily Burn Rate Telemetry", "Real-time burn velocity tracking (₹/day) against defined minimum safety reserves."),
            ("Insolvency Cliff Detection", "Proactive AI warnings identifying exact calendar dates where cash breaches safety thresholds.")
        ]),
        ("3.3 Invoices & Machine Learning Delay Radar", [
            ("Multi-Format Ingestion", "Drag-and-drop support for CSV, Excel (.xlsx/.xls), PDF (AI OCR scan), JSON (GST e-Invoice), and plain text."),
            ("ML Delay Prediction", "Trained model predicting exact payment delays (in days) per customer based on payment history and tenure."),
            ("Debtor Risk Scoring", "Automated classification of invoices into LOW, MEDIUM, and HIGH default risk tiers."),
            ("MSMED Act Section 15 Compliance", "Automated 45-day statutory clock tracking with automated Samadhaan recovery alert drafts.")
        ]),
        ("3.4 What-If Shock Simulator", [
            ("Interactive Stress Testing", "Adjustable sliders for Revenue Change (-50% to +50%), Expense Inflation, Payment Delays (0 to 60 days), Collection Efficiency, and GST Slabs."),
            ("Synthetic Twin vs. Baseline", "Side-by-side comparative Area charts illustrating baseline trajectory vs. simulated scenario."),
            ("Pre-Calibrated MSME Presets", "1-click simulations for 'Optimistic Growth', 'Severe Macro Shock', 'Payment Stagnation', and 'Base Case'.")
        ]),
        ("3.5 Risk Intelligence & Debtor Concentration", [
            ("Solvency Pressure Index", "Composite risk evaluation highlighting critical vulnerabilities before cash constraints occur."),
            ("Single-Buyer Concentration Radar", "Analysis of debtor concentration percentage to prevent existential dependency on key clients."),
            ("Actionable Mitigation Playbooks", "Prescriptive steps (e.g. TReDS factoring, vendor renegotiation, deposit reallocation) with quantified cash impacts.")
        ]),
        ("3.6 Expenses & Burn Optimization", [
            ("Itemized Expense Ledger", "Classification of operational expenditure into fixed recurring (payroll, rent, SaaS) and variable spend."),
            ("Category Distribution Donut", "Visual breakdown across Payroll, Facilities, Utilities, Logistics, Raw Materials, and Maintenance."),
            ("AI Burn Rationalization", "Targeted recommendations for vendor renegotiation, annual license conversion, and logistics consolidation.")
        ]),
        ("3.7 GST Intelligence & Statutory Reconciliation", [
            ("Live GSTR-3B Tax Computation", "Automated calculation of Output GST on Sales (GSTR-1), Input Tax Credit (ITC) from Purchases (GSTR-2B), and Net Cash Payable."),
            ("Interactive Multi-Slab Calculator", "Real-time calculation of CGST, SGST, IGST across 0%, 5%, 12%, 18%, and 28% slabs with inclusive/exclusive pricing toggles."),
            ("GSTIN Checksum Validator", "Instant structural and state-code validation for 15-character Indian GSTINs.")
        ]),
        ("3.8 Financing & TReDS Marketplace", [
            ("Non-Debt Liquidity Router", "Automated matching of verified invoices to TReDS (RXIL/M1xchange/Invoicemart) discounting at competitive APRs (7.8%–9.5%)."),
            ("Working Capital & Credit Lines", "Instant pre-qualification for collateral-free MSME credit lines and merchant cash advances."),
            ("Explainable Cost Comparison", "Clear breakdown of platform fees, interest rates, and total liquidity unlocked without hidden debt.")
        ]),
        ("3.9 Payroll & Worker Disbursal", [
            ("Staff & Wage Management", "Tracking factory workers, engineering staff, and contractors with attendance and role mappings."),
            ("Batch Salary Disbursal", "1-click salary processing with automated deductions, bonuses, and transaction reference generation."),
            ("Compliance Disbursal Statements", "Downloadable and printable statutory payroll audit logs.")
        ]),
        ("3.10 NexFin AI Copilot 2.0 & Multilingual Intelligence", [
            ("3-Pillar Root Cause Engine", "Interactive assistant answering complex financial queries with 'Root Cause', 'Immediate Impact', and 'Recommended Action'."),
            ("1-Click Action Triggers", "Direct execution buttons inside the chat drawer navigating to relevant diagnostic tools."),
            ("18 Indian Languages Support", "Complete UI localized in Hindi, Marathi, Gujarati, Tamil, Telugu, Kannada, Bengali, Punjabi, Malayalam, and more.")
        ])
    ]

    for mod_title, mod_features in modules:
        add_section_heading(mod_title, level=2)
        for feat_name, feat_desc in mod_features:
            add_bullet(None, feat_name, feat_desc)

    # -------------------------------------------------------------
    # 4. SECURITY, DATA PRIVACY & COMPLIANCE
    # -------------------------------------------------------------
    add_section_heading("4. Security, Data Privacy & Compliance Standards")
    add_bullet(None, "Client-Side Encryption & Storage", "User financial twins are persisted in authenticated user-partitioned local storage and encrypted backend datastores.")
    add_bullet(None, "Consent-Driven Architecture", "In accordance with India Account Aggregator (AA) and DPDP Act frameworks, data is only ingested with explicit user consent.")
    add_bullet(None, "Non-Custodial Financial Design", "NexFin acts purely as an analytical decision-support twin and does not hold or custody user funds.")
    add_bullet(None, "Audit Trail & Exportability", "All calculations and ledgers can be exported in standardized CSV format for statutory audits.")

    # Save document
    output_path = "d:/FRontend fix/FinTwin/NexFin_Platform_Documentation.docx"
    doc.save(output_path)
    print(f"Successfully generated documentation at: {output_path}")

if __name__ == "__main__":
    create_documentation()
